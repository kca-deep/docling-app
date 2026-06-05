"""
쇼케이스 추출 테스트용 샘플 문서 생성기

sample-showcase.md(원본 소스) 한 벌의 내용을 PDF/DOCX 두 포맷으로 동일하게 생성한다.
세 포맷(hwpx/pdf/docx)이 같은 콘텐츠를 갖도록 맞춰, kordoc 포맷별 추출 품질을
공정하게 비교/검증하기 위함이다. (hwpx는 별도 제작본을 그대로 사용)

실행:
    PYTHONPATH=. backend/venv/Scripts/python.exe backend/test_data/generate_samples.py

의존성: python-docx, reportlab (둘 다 backend/requirements 환경에 설치되어 있음)
한글 폰트: 맑은 고딕(C:/Windows/Fonts/malgun.ttf) 사용.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

BASE = Path(__file__).parent
SOURCE_MD = BASE / "sample-showcase.md"
OUT_PDF = BASE / "sample-showcase.pdf"
OUT_DOCX = BASE / "sample-showcase.docx"

KOREAN_FONT_PATH = Path("C:/Windows/Fonts/malgun.ttf")
KOREAN_FONT = "MalgunGothic"


def parse_blocks(md: str):
    """마크다운을 (종류, 텍스트) 블록 리스트로 단순 파싱.

    종류: h1/h2/h3, bullet, number, quote, para
    """
    blocks = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if m := re.match(r"^(#{1,6})\s+(.*)$", line):
            level = len(m.group(1))
            blocks.append((f"h{min(level, 3)}", m.group(2).strip()))
        elif m := re.match(r"^\d+[.)]\s+(.*)$", line):
            blocks.append(("number", m.group(1).strip()))
        elif m := re.match(r"^[-*+]\s+(.*)$", line):
            blocks.append(("bullet", m.group(1).strip()))
        elif m := re.match(r"^>\s+(.*)$", line):
            blocks.append(("quote", m.group(1).strip()))
        else:
            blocks.append(("para", line.strip()))
    return blocks


def build_docx(blocks):
    doc = Document()
    # 기본 폰트(한글) 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    num_idx = 0
    for kind, text in blocks:
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        elif kind == "quote":
            p = doc.add_paragraph(text)
            p.style = doc.styles["Intense Quote"] if "Intense Quote" in [s.name for s in doc.styles] else doc.styles["Normal"]
            for r in p.runs:
                r.italic = True
        else:
            doc.add_paragraph(text)
    doc.save(str(OUT_DOCX))


def build_pdf(blocks):
    pdfmetrics.registerFont(TTFont(KOREAN_FONT, str(KOREAN_FONT_PATH)))

    def style(name, size, **kw):
        return ParagraphStyle(
            name, fontName=KOREAN_FONT, fontSize=size, leading=size * 1.5,
            alignment=TA_LEFT, **kw,
        )

    styles = {
        "h1": style("h1", 18, spaceBefore=6, spaceAfter=10),
        "h2": style("h2", 14, spaceBefore=10, spaceAfter=6),
        "h3": style("h3", 12, spaceBefore=8, spaceAfter=4),
        "para": style("para", 11, spaceAfter=4),
        "bullet": style("bullet", 11, leftIndent=12, spaceAfter=2),
        "number": style("number", 11, leftIndent=12, spaceAfter=2),
        "quote": style("quote", 11, leftIndent=16, textColor="#444444", spaceBefore=2, spaceAfter=6),
    }

    flow = []
    num_idx = 0
    prev_kind = None
    for kind, text in blocks:
        if kind == "number":
            num_idx = num_idx + 1 if prev_kind == "number" else 1
            prefix = f"{num_idx}. "
        else:
            prefix = "• " if kind == "bullet" else ""
            if kind != "number":
                num_idx = 0
        # XML 이스케이프
        safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        flow.append(Paragraph(prefix + safe, styles.get(kind, styles["para"])))
        prev_kind = kind

    doc = SimpleDocTemplate(
        str(OUT_PDF), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=20 * mm, bottomMargin=20 * mm,
    )
    doc.build(flow)


def main():
    md = SOURCE_MD.read_text(encoding="utf-8")
    blocks = parse_blocks(md)
    build_docx(blocks)
    build_pdf(blocks)
    print(f"생성 완료: {OUT_DOCX.name}, {OUT_PDF.name} (blocks={len(blocks)})")


if __name__ == "__main__":
    main()
