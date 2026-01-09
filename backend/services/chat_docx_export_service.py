"""
채팅 DOCX 내보내기 서비스
LLM 응답 데이터를 Word 문서로 변환합니다.
"""

import logging
import re
from io import BytesIO
from typing import List, Tuple, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from backend.services.tool_executor_service import (
    ToolResult,
    file_storage,
    FileSizeExceededError,
    StorageCapacityExceededError
)

logger = logging.getLogger("uvicorn")


class ChatDocxExportService:
    """
    채팅 데이터를 Word 문서로 내보내는 서비스
    마크다운 형식의 텍스트를 지원합니다.
    """

    def __init__(self):
        pass

    def _parse_markdown_line(self, line: str) -> Tuple[str, str, int]:
        """
        마크다운 라인 파싱

        Returns:
            Tuple[type, content, level]
            - type: "heading", "bullet", "numbered", "text"
            - content: 실제 텍스트
            - level: 헤딩 레벨 또는 리스트 깊이
        """
        line = line.rstrip()

        # 헤딩 감지
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            return "heading", heading_match.group(2), level

        # 불릿 리스트 감지
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            level = indent // 2 + 1
            return "bullet", bullet_match.group(2), level

        # 번호 리스트 감지
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if numbered_match:
            indent = len(numbered_match.group(1))
            level = indent // 2 + 1
            return "numbered", numbered_match.group(2), level

        # 일반 텍스트
        return "text", line, 0

    def _add_formatted_text(self, paragraph, text: str):
        """
        인라인 마크다운 포맷팅 적용 (볼드, 이탤릭, 코드)
        """
        # 패턴: **bold**, *italic*, `code`
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'

        last_end = 0
        for match in re.finditer(pattern, text):
            # 매치 전 일반 텍스트
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            full_match = match.group(0)

            if full_match.startswith('**'):
                # Bold
                run = paragraph.add_run(match.group(2))
                run.bold = True
            elif full_match.startswith('*'):
                # Italic
                run = paragraph.add_run(match.group(3))
                run.italic = True
            elif full_match.startswith('`'):
                # Code
                run = paragraph.add_run(match.group(4))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)

            last_end = match.end()

        # 나머지 텍스트
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _parse_markdown_table(self, lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
        """
        마크다운 테이블 파싱

        Returns:
            Tuple[rows, end_idx]
        """
        rows = []
        idx = start_idx

        while idx < len(lines):
            line = lines[idx].strip()
            if not line or not line.startswith('|'):
                break

            # 구분선 스킵
            if re.match(r'^\|[\s\-:]+\|$', line) or line.replace('|', '').replace('-', '').replace(':', '').strip() == '':
                idx += 1
                continue

            # 셀 파싱
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for i, c in enumerate(cells) if c or i not in [0, len(cells)-1]]
            if cells:
                rows.append(cells)

            idx += 1

        return rows, idx

    def export_to_docx(
        self,
        content: str,
        title: str = "문서",
        filename: str = "document"
    ) -> bytes:
        """
        마크다운 텍스트를 Word 문서로 변환

        Args:
            content: 마크다운 형식의 텍스트
            title: 문서 제목
            filename: 파일명 (확장자 제외)

        Returns:
            bytes: Word 문서 바이너리
        """
        document = Document()

        # 제목 추가
        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 생성 날짜 추가
        date_para = document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%Y년 %m월 %d일"))
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        document.add_paragraph()  # 빈 줄

        # 본문 파싱 및 추가
        lines = content.split('\n')
        idx = 0

        while idx < len(lines):
            line = lines[idx]

            # 빈 줄
            if not line.strip():
                document.add_paragraph()
                idx += 1
                continue

            # 테이블 감지
            if line.strip().startswith('|'):
                rows, end_idx = self._parse_markdown_table(lines, idx)
                if rows:
                    self._add_table(document, rows)
                idx = end_idx
                continue

            # 코드 블록 감지
            if line.strip().startswith('```'):
                code_lines = []
                idx += 1
                while idx < len(lines) and not lines[idx].strip().startswith('```'):
                    code_lines.append(lines[idx])
                    idx += 1
                idx += 1  # closing ```

                if code_lines:
                    code_para = document.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                continue

            # 일반 라인 파싱
            line_type, text, level = self._parse_markdown_line(line)

            if line_type == "heading":
                doc_level = min(level, 4)  # Word는 최대 9레벨이지만 4까지만 사용
                document.add_heading(text, level=doc_level)

            elif line_type == "bullet":
                para = document.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)

            elif line_type == "numbered":
                para = document.add_paragraph(style='List Number')
                self._add_formatted_text(para, text)

            else:
                # 일반 텍스트
                if text.strip():
                    para = document.add_paragraph()
                    self._add_formatted_text(para, text)

            idx += 1

        # 바이너리로 저장
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _add_table(self, document: Document, rows: List[List[str]]):
        """Word 문서에 테이블 추가"""
        if not rows:
            return

        col_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=col_count)
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < col_count:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_value

                    # 첫 번째 행은 헤더로 스타일링
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        document.add_paragraph()  # 테이블 후 빈 줄

    async def handle_export_to_docx(
        self,
        tool_call_id: str,
        arguments: dict
    ) -> ToolResult:
        """
        export_to_docx 도구 핸들러

        Args:
            tool_call_id: 도구 호출 ID
            arguments: 도구 인자
                - content: 문서 내용
                - title: 문서 제목 (선택)
                - filename: 파일명 (선택)

        Returns:
            ToolResult: 실행 결과
        """
        try:
            content = arguments.get("content", "")
            if not content:
                return ToolResult(
                    tool_call_id=tool_call_id,
                    tool_name="export_to_docx",
                    success=False,
                    action_type="message",
                    error="문서 내용이 없습니다."
                )

            title = arguments.get("title", "문서")
            filename = arguments.get("filename", "document")

            # DOCX 생성
            docx_bytes = self.export_to_docx(content, title, filename)

            # 파일 저장소에 저장 (중복 확장자 방지)
            if filename.lower().endswith('.docx'):
                full_filename = filename
            else:
                full_filename = f"{filename}.docx"
            file_id = file_storage.store(
                filename=full_filename,
                content=docx_bytes,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            return ToolResult(
                tool_call_id=tool_call_id,
                tool_name="export_to_docx",
                success=True,
                action_type="download",
                file_id=file_id,
                filename=full_filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                message=f"Word 문서 '{full_filename}'이(가) 생성되었습니다."
            )

        except FileSizeExceededError as e:
            logger.warning(f"[DocxExport] File size exceeded: {e}")
            return ToolResult(
                tool_call_id=tool_call_id,
                tool_name="export_to_docx",
                success=False,
                action_type="message",
                error=str(e)
            )

        except StorageCapacityExceededError as e:
            logger.warning(f"[DocxExport] Storage capacity exceeded: {e}")
            return ToolResult(
                tool_call_id=tool_call_id,
                tool_name="export_to_docx",
                success=False,
                action_type="message",
                error=str(e)
            )

        except Exception as e:
            logger.error(f"[DocxExport] Export failed: {e}")
            return ToolResult(
                tool_call_id=tool_call_id,
                tool_name="export_to_docx",
                success=False,
                action_type="message",
                error=f"Word 문서 생성 중 오류가 발생했습니다: {str(e)}"
            )


# 싱글톤 인스턴스
chat_docx_export_service = ChatDocxExportService()
